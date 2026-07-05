import os
import glob
from pathlib import Path

# Importers
import extract_msg
import openpyxl
from docx import Document
from PyPDF2 import PdfReader

def extract_pdf(path):
    text = ""
    try:
        reader = PdfReader(path)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    except Exception as e:
        text = f"Error reading PDF: {e}"
    return text

def extract_docx(path):
    text = ""
    try:
        doc = Document(path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
    except Exception as e:
        text = f"Error reading DOCX: {e}"
    return text

def extract_xlsx(path):
    text = ""
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
        for sheetname in wb.sheetnames:
            ws = wb[sheetname]
            text += f"\n--- Sheet: {sheetname} ---\n"
            for row in ws.iter_rows(values_only=True):
                if any(row):
                    text += " | ".join([str(c) if c is not None else "" for c in row]) + "\n"
    except Exception as e:
        text = f"Error reading XLSX: {e}"
    return text

def extract_msg_file(path):
    text = ""
    try:
        msg = extract_msg.Message(path)
        text += f"From: {msg.sender}\n"
        text += f"To: {msg.to}\n"
        text += f"Date: {msg.date}\n"
        text += f"Subject: {msg.subject}\n"
        text += "-" * 20 + "\n"
        text += msg.body or ""
    except Exception as e:
        text = f"Error reading MSG: {e}"
    return text

def main():
    target_dir = r"d:\AntiGravity_Workspace\apps\ysdms-nextgen\source_data\SMK-230"
    output_file = os.path.join(target_dir, "extracted_content.md")
    
    files = glob.glob(os.path.join(target_dir, "*"))
    files.sort()
    
    with open(output_file, "w", encoding="utf-8") as f:
        f.write("# Extracted Content from SMK-230 Directory\n\n")
        
        for path in files:
            if os.path.basename(path).startswith("~$"):
                continue
            if path == output_file:
                continue
                
            ext = os.path.splitext(path)[1].lower()
            f.write(f"\n## File: {os.path.basename(path)}\n")
            f.write("```\n")
            
            if ext == ".pdf":
                f.write(extract_pdf(path))
            elif ext == ".docx" or ext == ".doc":
                if ext == ".doc":
                    f.write("Legacy .doc not fully supported by python-docx. Skipping or partial.\n")
                else:
                    f.write(extract_docx(path))
            elif ext == ".xlsx":
                f.write(extract_xlsx(path))
            elif ext == ".msg":
                f.write(extract_msg_file(path))
            elif ext in [".jpg", ".png", ".jpeg"]:
                f.write("Image file. Skipping content extraction.\n")
            else:
                f.write(f"Unsupported file type: {ext}\n")
                
            f.write("\n```\n")
            
    print(f"Extraction complete. Output saved to {output_file}")

if __name__ == "__main__":
    main()
